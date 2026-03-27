"""
Sensitivity classifier service.
Layer 1: PII regex detector (fast, high-precision) — runs per-line for accurate XAI
Layer 2: ML classifier (context-aware) — uses ML service endpoint if available, else local model
Layer 3: Decision logic (max of both layers if PII >= confidential)
"""

from __future__ import annotations
import csv
import io
import os
import re
import unicodedata
import logging
import httpx
from typing import Optional

from app.ml import features as feat_module
from app.ml import model as ml_model
from app.ml.features import (
    detect_pii_level,
    PII_PATTERNS,
    LEVEL_TO_INT,
    INT_TO_LEVEL,
    CONFIDENTIAL_IDX,
)

logger = logging.getLogger(__name__)

# ML Service endpoint (if configured, uses DistilBERT service; otherwise falls back to local model)
ML_SERVICE_URL = os.environ.get("AZURE_ML_ENDPOINT_URL", "")

# File processing imports (lazy to avoid import errors if libs not installed)
SUPPORT_PDF = True
SUPPORT_DOCX = True
try:
    import pdfplumber
except ImportError:
    SUPPORT_PDF = False
try:
    from docx import Document as DocxDocument
except ImportError:
    SUPPORT_DOCX = False

MAX_TEXT_LENGTH = 50_000
MAX_LINES_PER_SEGMENT = 5  # Lines grouped per segment for per-segment display
SUPPORTED_MIMES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "text/markdown",
    "text/x-markdown",
    "text/csv",
    "application/csv",
}
FILE_SIGNATURES = {
    "application/pdf": b"%PDF-",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": b"PK",
}

# Markdown syntax to strip before classification
_MD_SYNTAX = re.compile(
    r"(`{1,3}[^`]*`{1,3})"  # inline/block code
    r"|(!?\[.*?\]\(.*?\))"  # images & links
    r"|^#{1,6}\s"  # headings
    r"|^[-*_]{3,}$"  # horizontal rules
    r"|\*{1,2}([^*]+)\*{1,2}"  # bold/italic
    r"|_{1,2}([^_]+)_{1,2}",  # underscores
    re.MULTILINE,
)


def extract_text(file_bytes: bytes, content_type: str) -> str:
    """Extract plain text from PDF, DOCX, TXT, MD, or CSV bytes."""
    ct = content_type.split(";")[0].strip().lower()
    if ct == "application/pdf":
        if not SUPPORT_PDF:
            raise ValueError("pdfplumber not installed")
        return _extract_pdf(file_bytes)
    elif (
        ct == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        if not SUPPORT_DOCX:
            raise ValueError("python-docx not installed")
        return _extract_docx(file_bytes)
    elif ct in ("text/plain",):
        return _decode_text(file_bytes)
    elif ct in ("text/markdown", "text/x-markdown"):
        return _extract_markdown(file_bytes)
    elif ct in ("text/csv", "application/csv"):
        return _extract_csv(file_bytes)
    else:
        raise ValueError(f"Unsupported content type: {content_type}")


def _decode_text(data: bytes) -> str:
    """Safe UTF-8 decode with NFC normalization."""
    text = data.decode("utf-8", errors="replace")
    return unicodedata.normalize("NFC", text)[:MAX_TEXT_LENGTH]


def _extract_pdf(file_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        if len(pdf.pages) > 200:
            raise ValueError("PDF exceeds 200-page limit")
        for page in pdf.pages:
            # Extract text with layout detection for better accuracy
            page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if page_text:
                text_parts.append(page_text.strip())
            # Also extract table cells
            for table in page.extract_tables():
                for row in table:
                    if row:
                        row_text = " | ".join(cell or "" for cell in row if cell)
                        if row_text.strip():
                            text_parts.append(row_text)
    full = "\n\n".join(text_parts)
    return full[:MAX_TEXT_LENGTH]


def _extract_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Extract table cells too
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
            if row_text.strip():
                parts.append(row_text)
    text = "\n".join(parts)
    return text[:MAX_TEXT_LENGTH]


def _extract_markdown(file_bytes: bytes) -> str:
    """Strip markdown syntax and return plain text for classification."""
    raw = _decode_text(file_bytes)
    # Remove code fences (multiline)
    raw = re.sub(r"```[\s\S]*?```", "", raw)
    raw = re.sub(r"~~~[\s\S]*?~~~", "", raw)
    # Strip inline markup
    cleaned = _MD_SYNTAX.sub(lambda m: m.group(3) or m.group(4) or "", raw)
    # Remove HTML tags if any
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    # Collapse multiple blank lines
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()[:MAX_TEXT_LENGTH]


def _extract_csv(file_bytes: bytes) -> str:
    """Flatten CSV into text: header names + all cell values."""
    raw = file_bytes.decode("utf-8", errors="replace")
    parts = []
    try:
        reader = csv.reader(io.StringIO(raw))
        rows = list(reader)
        if rows:
            parts.append("Headers: " + ", ".join(rows[0]))
        for row in rows[1:]:
            parts.append(" | ".join(cell for cell in row if cell.strip()))
    except Exception:
        # Fallback: treat as plain text
        parts.append(raw)
    return "\n".join(parts)[:MAX_TEXT_LENGTH]


def validate_file_magic(file_bytes: bytes, content_type: str) -> None:
    """Magic byte validation to prevent MIME spoofing."""
    ct = content_type.split(";")[0].strip()
    expected = FILE_SIGNATURES.get(ct)
    if expected and not file_bytes.startswith(expected):
        raise ValueError(f"File content does not match declared type: {content_type}")


# ─── PII Match Helpers ────────────────────────────────────────────────────────

_PATTERN_META = [
    ("ssn", PII_PATTERNS["ssn"], "highly_sensitive", 0.42, "Social Security Number"),
    (
        "credit_card",
        PII_PATTERNS["credit_card"],
        "highly_sensitive",
        0.45,
        "Credit/Debit Card Number",
    ),
    (
        "password_kw",
        PII_PATTERNS["password_kw"],
        "highly_sensitive",
        0.50,
        "Password or API Key",
    ),
    ("aadhaar", PII_PATTERNS["aadhaar"], "highly_sensitive", 0.40, "Aadhaar Number"),
    ("email", PII_PATTERNS["email"], "confidential", 0.18, "Email Address"),
    ("phone", PII_PATTERNS["phone"], "confidential", 0.15, "Phone Number"),
    ("pan", PII_PATTERNS["pan"], "confidential", 0.35, "PAN Number"),
    ("dob", PII_PATTERNS["dob"], "confidential", 0.30, "Date of Birth"),
    ("ip_address", PII_PATTERNS["ip_address"], "internal", 0.10, "IP Address"),
]

_SENSITIVE_KW_META = [
    ("salary", "highly_sensitive", 0.35, "Salary/Compensation Data"),
    ("diagnosis", "highly_sensitive", 0.35, "Medical/Diagnosis Data"),
    ("prescription", "highly_sensitive", 0.30, "Medical Prescription"),
    ("patient", "confidential", 0.20, "Patient Information"),
    ("confidential", "confidential", 0.20, "Confidential Marker"),
    ("classified", "confidential", 0.20, "Classified Marker"),
    ("medical", "confidential", 0.18, "Medical Information"),
    ("account number", "highly_sensitive", 0.38, "Bank Account Number"),
    ("routing", "highly_sensitive", 0.35, "Routing Number"),
    ("credential", "highly_sensitive", 0.40, "Credential Data"),
    ("merger", "confidential", 0.22, "M&A Activity"),
    ("acquisition", "confidential", 0.22, "M&A Activity"),
    ("settlement", "confidential", 0.20, "Legal Settlement"),
    ("termination", "confidential", 0.18, "HR/Termination Record"),
    ("compensation", "confidential", 0.18, "Compensation Data"),
    ("bonus", "confidential", 0.15, "Compensation Data"),
    ("private", "internal", 0.10, "Internal/Private Marker"),
    ("internal", "internal", 0.08, "Internal Document Marker"),
]


def _scan_line_for_pii(line: str, line_num: int) -> list[dict]:
    """
    Return a list of real PII match findings for a single line.
    Each entry contains the pattern name, exact match text, column offset, and level.
    No data is fabricated — every entry corresponds to an actual regex match.
    """
    findings = []
    for pname, pattern, level, weight, label in _PATTERN_META:
        for m in pattern.finditer(line):
            findings.append(
                {
                    "pattern": pname,
                    "label": label,
                    "match": m.group(0),
                    "line": line_num,
                    "col_start": m.start() + 1,  # 1-indexed
                    "col_end": m.end() + 1,
                    "level": level,
                    "weight": weight,
                }
            )
    # Keyword scan
    line_lower = line.lower()
    for kw, kw_level, kw_weight, kw_label in _SENSITIVE_KW_META:
        if kw in line_lower:
            idx = line_lower.index(kw)
            findings.append(
                {
                    "pattern": "keyword",
                    "label": kw_label,
                    "match": kw,
                    "line": line_num,
                    "col_start": idx + 1,
                    "col_end": idx + len(kw) + 1,
                    "level": kw_level,
                    "weight": kw_weight,
                }
            )
            break  # One keyword per line is enough info
    return findings


def _level_display(level: str) -> str:
    return level.replace("_", " ").title()


def _call_ml_service(text: str) -> tuple[str, float, str]:
    """
    Call the ML service endpoint for classification.

    Returns: (level, confidence, model_version)
    Falls back to local model if ML service is unavailable.
    """
    if not ML_SERVICE_URL:
        # No ML service configured, use local model
        ml_level, ml_confidence = ml_model.predict(text)
        return ml_level, ml_confidence, ml_model.get_model_version()

    try:
        # Call ML service endpoint
        with httpx.Client(timeout=10.0) as client:
            response = client.post(
                ML_SERVICE_URL, json={"texts": [text], "include_explanation": True}
            )
            response.raise_for_status()

        data = response.json()
        results = data.get("results", [])

        if results:
            result = results[0]
            level = result.get("sensitivity_level", "internal")
            confidence = result.get("confidence", 0.8)
            model_version = data.get("model_version", "distilbert-mnli-v1.0")
            logger.info(f"ML Service classification: {level} ({confidence:.2%})")
            return level, confidence, f"ml_service:{model_version}"

    except Exception as e:
        logger.warning(f"ML Service call failed, using local model: {e}")

    # Fallback to local model
    ml_level, ml_confidence = ml_model.predict(text)
    return ml_level, ml_confidence, ml_model.get_model_version()


def classify_text(text: str) -> dict:
    """
    Hybrid classification pipeline (fast version, no segments).
    Returns: {level, confidence, explanation_factors, explanation_summary, model_version}
    """
    text = text[:MAX_TEXT_LENGTH].replace("\x00", "")

    # Layer 1: PII rule-based
    pii_level_int = detect_pii_level(text)

    # Layer 2: ML classifier (uses ML service if available, else local model)
    ml_level, ml_confidence, version = _call_ml_service(text)
    ml_level_int = LEVEL_TO_INT.get(ml_level, 1)  # default to internal if unknown

    # Layer 3: Decision
    if pii_level_int >= CONFIDENTIAL_IDX:
        final_level_int = max(pii_level_int, ml_level_int)
        confidence = max(0.92, ml_confidence)
    else:
        final_level_int = ml_level_int
        confidence = ml_confidence

    final_level = INT_TO_LEVEL[final_level_int]
    factors = _build_explanation_factors(text)
    summary = _build_summary(final_level, confidence, factors)

    return {
        "level": final_level,
        "confidence": round(confidence, 4),
        "explanation_factors": factors,
        "explanation_summary": summary,
        "model_version": version,
    }


def classify_text_detailed(text: str, source_label: str = "text") -> dict:
    """
    Full classification with per-line/segment PII detail.

    Strategy:
    - Run ML model ONCE on full text → overall level & confidence (accuracy preserved)
    - Scan each line with PII regex → exact match locations, no fabrication
    - Segments group MAX_LINES_PER_SEGMENT lines for readable display

    Returns: {level, confidence, explanation_factors, explanation_summary,
              model_version, segments: [...]}
    """
    text = text[:MAX_TEXT_LENGTH].replace("\x00", "")
    lines = text.splitlines()
    if not lines:
        return classify_text(text)

    # ── ML runs once on full text (uses ML service if available) ─────────────────
    pii_level_int = detect_pii_level(text)
    ml_level, ml_confidence, version = _call_ml_service(text)
    ml_level_int = LEVEL_TO_INT.get(ml_level, 1)

    if pii_level_int >= CONFIDENTIAL_IDX:
        overall_level_int = max(pii_level_int, ml_level_int)
        confidence = max(0.92, ml_confidence)
    else:
        overall_level_int = ml_level_int
        confidence = ml_confidence

    overall_level = INT_TO_LEVEL[overall_level_int]

    # ── Per-line PII scan for accurate location-based XAI ─────────────────────
    all_findings: list[dict] = []  # (line_num, finding)
    findings_by_line: dict[int, list] = {}

    for i, line in enumerate(lines, start=1):
        if not line.strip():
            continue
        line_findings = _scan_line_for_pii(line, i)
        if line_findings:
            findings_by_line[i] = line_findings
            all_findings.extend(line_findings)

    # ── Build segments (groups of MAX_LINES_PER_SEGMENT lines) ───────────────
    segments = []
    chunk_size = MAX_LINES_PER_SEGMENT

    for chunk_start in range(0, len(lines), chunk_size):
        chunk_lines = lines[chunk_start : chunk_start + chunk_size]
        line_start = chunk_start + 1
        line_end = chunk_start + len(chunk_lines)

        # Collect findings for this chunk
        chunk_findings = []
        for ln in range(line_start, line_end + 1):
            chunk_findings.extend(findings_by_line.get(ln, []))

        # Determine segment level
        if chunk_findings:
            seg_level = max(
                chunk_findings, key=lambda f: LEVEL_TO_INT.get(f["level"], 0)
            )["level"]
        else:
            seg_level = "public"

        # Build human-readable reasons from actual matches
        reasons = []
        seen_patterns = set()
        for f in chunk_findings:
            key = f["pattern"] + f["match"][:20]
            if key not in seen_patterns:
                seen_patterns.add(key)
                reasons.append(
                    {
                        "pattern": f["pattern"],
                        "label": f["label"],
                        "match": f["match"][:60],  # Truncate long matches
                        "line": f["line"],
                        "col_start": f["col_start"],
                    }
                )

        # Content preview: first non-empty line of chunk, truncated
        preview_lines = [l for l in chunk_lines if l.strip()]
        content_preview = (
            (preview_lines[0][:120] + "…")
            if preview_lines and len(preview_lines[0]) > 120
            else (preview_lines[0] if preview_lines else "")
        )

        # Explanation text
        if reasons:
            reason_strs = []
            for r in reasons[:3]:
                reason_strs.append(
                    f"'{r['label']}' detected at line {r['line']}, col {r['col_start']}"
                )
            explanation = "; ".join(reason_strs) + "."
        else:
            explanation = (
                f"No sensitive patterns detected in lines {line_start}–{line_end}."
            )

        # Only include segments that have content
        if any(l.strip() for l in chunk_lines):
            segments.append(
                {
                    "segment_id": len(segments) + 1,
                    "source": source_label,
                    "line_start": line_start,
                    "line_end": line_end,
                    "content_preview": content_preview,
                    "level": seg_level,
                    "level_int": LEVEL_TO_INT.get(seg_level, 0),
                    "has_pii": bool(chunk_findings),
                    "reasons": reasons[:5],  # Cap at 5 per segment
                    "explanation": explanation,
                }
            )

    # Cap segments to avoid enormous responses
    if len(segments) > 100:
        # Keep all high-sensitivity segments, sample others
        high = [s for s in segments if s["level_int"] >= CONFIDENTIAL_IDX]
        low = [s for s in segments if s["level_int"] < CONFIDENTIAL_IDX]
        segments = high + low[: (100 - len(high))]
        segments.sort(key=lambda s: s["segment_id"])

    # ── Build overall explanation factors from all findings ─────────────────────
    factors = _build_explanation_factors_from_findings(all_findings, text)
    summary = _build_summary(overall_level, confidence, factors)

    return {
        "level": overall_level,
        "confidence": round(confidence, 4),
        "explanation_factors": factors,
        "explanation_summary": summary,
        "model_version": version,
        "segments": segments,
        "total_findings": len(all_findings),
    }


def classify_pdf_detailed(file_bytes: bytes) -> dict:
    """
    PDF-specific: classify per-page for better file-level XAI.
    Returns combined result with page-level segments.
    """
    if not SUPPORT_PDF:
        raise ValueError("pdfplumber not installed")

    all_segments = []
    all_texts = []
    segment_id = 1

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        if len(pdf.pages) > 200:
            raise ValueError("PDF exceeds 200-page limit")

        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text(x_tolerance=3, y_tolerance=3) or ""
            # Also get tables
            for table in page.extract_tables():
                for row in table:
                    if row:
                        row_text = " | ".join(cell or "" for cell in row if cell)
                        if row_text.strip():
                            page_text += "\n" + row_text

            if not page_text.strip():
                continue

            all_texts.append(page_text)

            # Per-page PII scan
            page_lines = page_text.splitlines()
            page_findings = []
            for ln_idx, line in enumerate(page_lines, start=1):
                if line.strip():
                    page_findings.extend(_scan_line_for_pii(line, ln_idx))

            if page_findings:
                seg_level = max(
                    page_findings, key=lambda f: LEVEL_TO_INT.get(f["level"], 0)
                )["level"]
            else:
                seg_level = "public"

            reasons = []
            seen = set()
            for f in page_findings:
                key = f["pattern"] + f["match"][:20]
                if key not in seen:
                    seen.add(key)
                    reasons.append(
                        {
                            "pattern": f["pattern"],
                            "label": f["label"],
                            "match": f["match"][:60],
                            "line": f["line"],
                            "col_start": f["col_start"],
                        }
                    )

            preview = page_text[:100].replace("\n", " ").strip()
            if reasons:
                explanation = (
                    "; ".join(
                        f"'{r['label']}' at page {page_num}, line {r['line']}"
                        for r in reasons[:3]
                    )
                    + "."
                )
            else:
                explanation = f"No sensitive patterns on page {page_num}."

            all_segments.append(
                {
                    "segment_id": segment_id,
                    "source": "page",
                    "page": page_num,
                    "line_start": 1,
                    "line_end": len(page_lines),
                    "content_preview": (preview[:120] + "…")
                    if len(preview) > 120
                    else preview,
                    "level": seg_level,
                    "level_int": LEVEL_TO_INT.get(seg_level, 0),
                    "has_pii": bool(page_findings),
                    "reasons": reasons[:5],
                    "explanation": explanation,
                }
            )
            segment_id += 1

    full_text = "\n\n".join(all_texts)[:MAX_TEXT_LENGTH]
    if not full_text.strip():
        return {
            "level": "public",
            "confidence": 0.5,
            "explanation_factors": [],
            "explanation_summary": "No text content found in PDF.",
            "segments": [],
        }

    pii_level_int = detect_pii_level(full_text)
    ml_level, ml_confidence = ml_model.predict(full_text)
    ml_level_int = LEVEL_TO_INT[ml_level]

    if pii_level_int >= CONFIDENTIAL_IDX:
        overall_level_int = max(pii_level_int, ml_level_int)
        confidence = max(0.92, ml_confidence)
    else:
        overall_level_int = ml_level_int
        confidence = ml_confidence

    overall_level = INT_TO_LEVEL[overall_level_int]
    all_findings = [r for seg in all_segments for r in seg["reasons"]]
    factors = _build_explanation_factors_from_findings(all_findings, full_text)
    summary = _build_summary(overall_level, confidence, factors)

    return {
        "level": overall_level,
        "confidence": round(confidence, 4),
        "explanation_factors": factors,
        "explanation_summary": summary,
        "model_version": ml_model.get_model_version(),
        "segments": all_segments,
        "total_findings": len(all_findings),
        "extracted_text": full_text,
    }


def _build_explanation_factors_from_findings(
    findings: list[dict], text: str
) -> list[dict]:
    """Build explanation factors from actual PII match findings — no fabrication."""
    seen_patterns: dict[str, dict] = {}

    for f in findings:
        pname = f["pattern"]
        if pname not in seen_patterns:
            seen_patterns[pname] = {
                "feature": f"{f['label']} ({pname}) detected — {f['match'][:40]}",
                "weight": f["weight"],
                "count": 1,
            }
        else:
            seen_patterns[pname]["count"] += 1
            seen_patterns[pname]["feature"] = (
                f"{f['label']} ({pname}) detected — {seen_patterns[pname]['count']} occurrences"
            )

    factors = list(seen_patterns.values())

    # Add ML-based features if no PII found
    if not factors:
        text_lower = text.lower()
        for kw, kw_level, kw_weight, kw_label in _SENSITIVE_KW_META:
            if kw in text_lower:
                factors.append(
                    {
                        "feature": f"Sensitive keyword: '{kw}' ({kw_label})",
                        "weight": kw_weight,
                    }
                )
                if len(factors) >= 6:
                    break

    factors.sort(key=lambda x: x["weight"], reverse=True)
    return [
        {"feature": f["feature"], "weight": round(f["weight"], 3)} for f in factors[:6]
    ]


def _build_explanation_factors(text: str) -> list[dict]:
    """Legacy explanation builder for the simple classify_text() path."""
    findings = []
    for ln_idx, line in enumerate(text.splitlines()[:200], start=1):  # Cap at 200 lines
        if line.strip():
            findings.extend(_scan_line_for_pii(line, ln_idx))
    return _build_explanation_factors_from_findings(findings, text)


def _build_summary(level: str, confidence: float, factors: list[dict]) -> str:
    level_display = _level_display(level)
    pct = round(confidence * 100, 1)
    if factors:
        top = "; ".join(f["feature"] for f in factors[:3])
        return f"Classified as {level_display} (confidence: {pct}%) based on: {top}."
    return f"Classified as {level_display} (confidence: {pct}%) based on statistical text patterns."
